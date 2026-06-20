"""
Aerumentis — Document Processor
Extracts text from PDFs, DOCX, and plain text. Intelligent chunking with overlap.
"""
from __future__ import annotations

import hashlib
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import tiktoken

from aerumentis.core.config import get_settings
from aerumentis.core.logging import get_logger

logger = get_logger("aerumentis.document_processor")
settings = get_settings()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}


@dataclass
class DocumentChunk:
    id: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    token_count: int = 0

    def __post_init__(self) -> None:
        if self.token_count == 0:
            self.token_count = count_tokens(self.text)


@dataclass
class ProcessedDocument:
    id: str
    filename: str
    file_type: str
    total_text: str
    chunks: list[DocumentChunk] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    checksum: str = ""

    def __post_init__(self) -> None:
        if not self.checksum:
            self.checksum = hashlib.sha256(self.total_text.encode()).hexdigest()[:16]


_encoder: tiktoken.Encoding | None = None


def count_tokens(text: str) -> int:
    global _encoder
    if _encoder is None:
        try:
            _encoder = tiktoken.get_encoding("cl100k_base")
        except Exception:
            return len(text) // 4
    return len(_encoder.encode(text))


def extract_text_from_pdf(file_path: str) -> tuple[str, dict[str, Any]]:
    import pdfplumber
    pages: list[str] = []
    metadata: dict[str, Any] = {"page_count": 0}
    with pdfplumber.open(file_path) as pdf:
        metadata["page_count"] = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text:
                pages.append(f"--- Page {i + 1} ---\n{text}")
            tables = page.extract_tables()
            if tables:
                for j, table in enumerate(tables):
                    table_text = "\n".join(" | ".join(str(cell or "") for cell in row) for row in table)
                    pages.append(f"--- Page {i + 1}, Table {j + 1} ---\n{table_text}")
    return "\n\n".join(pages), metadata


def extract_text_from_docx(file_path: str) -> tuple[str, dict[str, Any]]:
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), {"paragraph_count": len(paragraphs)}


def extract_text_from_txt(file_path: str) -> tuple[str, dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return text, {"char_count": len(text)}


def extract_text_from_html(file_path: str) -> tuple[str, dict[str, Any]]:
    from bs4 import BeautifulSoup
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        html = f.read()
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return text, {"title": soup.title.string if soup.title else ""}


EXTRACTORS = {
    ".pdf": extract_text_from_pdf, ".docx": extract_text_from_docx,
    ".txt": extract_text_from_txt, ".md": extract_text_from_txt,
    ".html": extract_text_from_html, ".htm": extract_text_from_html,
}


def extract_text(file_path: str) -> tuple[str, dict[str, Any]]:
    ext = Path(file_path).suffix.lower()
    extractor = EXTRACTORS.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")
    return extractor(file_path)


def chunk_text(
    text: str, chunk_size: int | None = None, overlap: int | None = None,
    separators: list[str] | None = None,
) -> list[str]:
    chunk_size = chunk_size or settings.rag_chunk_size
    overlap = overlap or settings.rag_chunk_overlap
    separators = separators or ["\n\n\n", "\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " "]

    def _split(text: str, seps: list[str]) -> list[str]:
        if len(text) <= chunk_size:
            return [text] if text.strip() else []
        if not seps:
            return _split_by_length(text, chunk_size, overlap)
        sep = seps[0]
        remaining = seps[1:]
        if sep == "":
            return _split_by_length(text, chunk_size, overlap)
        splits = text.split(sep)
        chunks: list[str] = []
        current = ""
        for part in splits:
            candidate = current + sep + part if current else part
            if len(candidate) <= chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                if len(part) > chunk_size:
                    chunks.extend(_split(part, remaining))
                    current = ""
                else:
                    current = part
        if current:
            chunks.append(current)
        return [c for c in chunks if c.strip()]

    def _split_by_length(text: str, size: int, overlap: int) -> list[str]:
        chunks = []
        start = 0
        while start < len(text):
            chunks.append(text[start : start + size])
            start = start + size - overlap
            if start >= len(text):
                break
        return chunks

    return _split(text, separators)


def process_document(
    file_path: str, filename: str | None = None,
    extra_metadata: dict[str, Any] | None = None,
    chunk_size: int | None = None, overlap: int | None = None,
) -> ProcessedDocument:
    file_path_obj = Path(file_path)
    filename = filename or file_path_obj.name
    ext = file_path_obj.suffix.lower()
    logger.info("processing_document", filename=filename, file_type=ext)

    text, extraction_metadata = extract_text(str(file_path_obj))
    if not text.strip():
        raise ValueError(f"No text content extracted from {filename}")

    chunk_texts = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    doc_id = str(uuid.uuid4())
    chunks: list[DocumentChunk] = []

    for i, chunk_content in enumerate(chunk_texts):
        chunk_metadata = {
            "document_id": doc_id, "filename": filename, "file_type": ext,
            "chunk_index": i, "total_chunks": len(chunk_texts),
            **extraction_metadata, **(extra_metadata or {}),
        }
        chunks.append(DocumentChunk(id=str(uuid.uuid4()), text=chunk_content, metadata=chunk_metadata))

    document = ProcessedDocument(
        id=doc_id, filename=filename, file_type=ext, total_text=text, chunks=chunks,
        metadata={**extraction_metadata, **(extra_metadata or {}),
                  "chunk_count": len(chunks), "total_tokens": sum(c.token_count for c in chunks)},
    )
    logger.info("document_processed", filename=filename, chunks=len(chunks),
                total_tokens=document.metadata["total_tokens"])
    return document
